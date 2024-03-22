import csv
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
def add_table_to_docx(csv_file_path, document):
    # Read the CSV file
    with open(csv_file_path, newline='') as f:
        csv_reader = csv.reader(f) 

        csv_headers = next(csv_reader)
        csv_cols = len(csv_headers)

        table = document.add_table(rows=2, cols=csv_cols,style="Table Grid")
        hdr_cells = table.rows[0].cells

        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(csv_cols):
                row_cells[i].text = row[i]
    document.add_paragraph()

    # Save the Word documentument

    