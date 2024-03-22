import os
from fastapi import HTTPException
import math
from fastapi.responses import RedirectResponse, FileResponse
from fastapi import HTTPException, status
from tempfile import NamedTemporaryFile
from pdf2image import convert_from_bytes,convert_from_path
import pytesseract
from fastapi import FastAPI, UploadFile, File
import uvicorn
from fastapi.middleware.cors import CORSMiddleware
import layoutparser as lp
import cv2
from bs4 import BeautifulSoup
import subprocess
import PyPDF2
from pdf2docx import parse
from paddleocr import PaddleOCR
import io
import shutil
import gc
from PIL import Image
import numpy as np
import pdfplumber
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import string
import csv
from tableExtractor import TableExtractionPipeline
import asyncio
app = FastAPI(
    max_request_size=100000000  # Adjust this value based on your requirements
)
table_pipeline = TableExtractionPipeline()

ocr = PaddleOCR(use_angle_cls=False, lang='en',show_log=False) 
model = lp.PaddleDetectionLayoutModel(config_path="lp://PubLayNet/ppyolov2_r50vd_dcn_365e_publaynet/config",
                                threshold=0.5,
                                label_map={0: "Text", 1: "Title", 2: "List", 3:"Table", 4:"Figure"},
                                enforce_cpu=False,
                                enable_mkldnn=True)

origins = [
    "http://localhost",
    "http://localhost:3000",
    "https://lingo-staging.netlify.app",
    "http://localhost:8000",
    "https://lingoyouniverseocr.com",
    "https://206.189.230.159/",
    "https://www.lingoyouniverse.com"
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
) 

def preprocess_image(image):
    # Convert PIL imge to OpenCV format
    open_cv_image = np.array(image)
    open_cv_image = open_cv_image[:, :, ::-1].copy()

    # Apply preprocessing techniques such as image binarization, noise reduction, and deskewing
    gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    _, threshold_img = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)

    # Convert back to PIL image format
    pil_image = Image.fromarray(threshold_img)
    return pil_image
def pytess_for_list(cell_pil_img):
    cell_np_img = np.array(cell_pil_img)
    results = ocr.ocr(cell_np_img, cls=False,det=True,rec=True)
    result_txt=[]
    for result in results:
        if result != None:
            for line in result:
                if line != None:
                    #print(len(line))
                    #print("text", line[1][0])
                    line_list = list(line[1])
                    if ']' in line_list[0] and '[' in line_list[0]:
                        result_txt.append("\n\n")
                    if ';' in line_list[0] or ':' in line_list[0]:
                        line_list[0] = line_list[0].replace(';', ';\n').replace(':', ':\n')
                    result_txt.append(line_list[0])
                    if len(result_txt)<=22:  
                        #print(len(result_txt))
                        result_txt.append("\n")                          
    return ' '.join(result_txt).strip()


def pytess_for_image(cell_pil_img):
    cell_np_img = np.array(cell_pil_img)
    results = ocr.ocr(cell_pil_img, cls=False,det=True,rec=True)
    result_txt=[]
    for result in results:
        if result != None:
            for line in result:
                if line != None:
                    #print(len(line))
                    #print("text", line[1][0])
                    result_txt.append(line[1][0]) 
                    result_txt.append("\n")          
    return ' '.join(result_txt).strip()

def is_scanned_pdf(pdf_file):
    text=""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text+=page.extract_text()
        if text=="":
            return True
        else:
            return False
         
def add_table_to_docx(csv_file_path, document):
    # Read the CSV file
    with open(csv_file_path, 'r') as file:
        reader = csv.reader(file)
        data = list(reader)

    # Create a new Word document
    

    # Add a table to the Word document
    table = document.add_table(rows=0, cols=0)
    for row_data in data:
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            cell.text = row_data[i]

    # Save the Word documentument
    
def add_figure_table_to_docx(document, image_path):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(4))  # Adjust the width as needed
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the figure to the left
def add_figure_to_docx(document, image_path):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(4))  # Adjust the width as needed
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the figure to the left

    # Perform OCR on the image
    #extracted_text = pytesseract.image_to_string(image_path)
    extracted_text=pytess_for_image(image_path)+'\n'
   # text_to_add=sanitize_text(extracted_text)
    # Add the extracted text below the image
    text_paragraph = document.add_paragraph()
    text_run = text_paragraph.add_run(extracted_text)
def add_title_to_docx(document, title_text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title_text)
    font = run.font
    font.name = 'Courier New'
    font.bold = True  # Set the font to bold
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Inches(5 / 72)  # Adjust the spacing before the title
    paragraph_format.space_after = Inches(5 / 72)  # Adjust the spacing after the title
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align the title to the center


def sanitize_text(text):
    """
    Sanitize the text to remove non-XML-compatible characters, leading, and trailing whitespaces or newline characters.
    """
    # Define the printable characters, which are considered XML-compatible.
    printable = set(string.printable)

    # Filter out non-XML-compatible characters.
    sanitized_text = ''.join(filter(lambda x: x in printable, text))

    # Remove leading and trailing whitespaces and newline characters
    sanitized_text = sanitized_text.strip()

    return sanitized_text

def add_list_to_docx(document, text):
   # print("inside text",print(text))
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Arial'
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Inches(2.5 / 72)
    paragraph_format.left_indent = Inches(4 / 72)
    paragraph_format.space_after = Inches(2 / 72)
    paragraph_format.right_indent = Inches(3.5 / 72)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 2.1
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Function to add text to a Word document
def add_text_to_docx(document, text):
   # print("inside text",print(text))
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Arial'
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Inches(2.5 / 72)
    paragraph_format.left_indent = Inches(4 / 72)
    paragraph_format.space_after = Inches(2 / 72)
    paragraph_format.right_indent = Inches(3.5 / 72)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 2.1
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

@app.get("/")
async def root():
    # redirect to /docs
    redirect_url = "/docs"
    return RedirectResponse(url=redirect_url)

@app.post("/upload/")
async def upload_file(pdf_file: UploadFile = File(...)):
    """The main endpoint that the request comes in

    Args:
        pdf_file (UploadFile, optional): _description_. Defaults to File(...).

    Raises:
        HTTPException: _description_
        HTTPException: _description_

    Returns:
        _type_: A docx file. Returns FastAPI FileResponse object with media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    """    
    try:
        with open(pdf_file.filename, "wb") as buffer:
            shutil.copyfileobj(pdf_file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error saving PDF file: {str(e)}")

    # Check if the PDF pages are scanned
    try:  
        if is_scanned_pdf(pdf_file.filename):
            # The above code is creating a directory for storing output images and output CSV files.
            # It then converts a PDF file to a series of JPEG images using the `convert_from_path`
            # function. It creates a Word document using the `Document` class from the `python-docx`
            # library. It calculates the total number of images and determines the batch size and
            # number of batches for processing the images.
            output_folder_img = 'images'
            output_directory = "output/output_csv"
            os.makedirs(output_directory, exist_ok=True)
            os.makedirs(output_folder_img, exist_ok=True)
            images = convert_from_path(pdf_file.filename, fmt='jpeg')
            docx_path = "output.docx" 
            document = Document()
            total_images = len(images)
            batch_size = 20
            num_batches = math.ceil(total_images / batch_size)

         # The above code is iterating over a set of images in batches. It calculates the start and
         # end indices for each batch, selects the images for the current batch, and then performs
         # some processing on each image.
            for batch_idx in range(num_batches):
                start_idx = batch_idx * batch_size
                end_idx = min(total_images, start_idx + batch_size)
                batch_images = images[start_idx:end_idx]

                for i, image in enumerate(batch_images):
                
                    page_number = i + 1  # Calculate the correct page number
                    #processed_image = preprocess_image(image) 
                    image_path = f'{output_folder_img}/output_page_{page_number}.png'
                    image.save(image_path, 'PNG')
                    img = cv2.imread(image_path)
                    img = img[..., ::-1]
                    layout = model.detect(img)
                    # The above code is creating a new `Layout` object called `text_blocks` by
                    # iterating over each element `b` in the `layout` object and adding it to a list
                    # comprehension.
                  
                    # The above code is performing the following tasks:
                    text_blocks = lp.Layout([b for b in layout])
                    h, w = img.shape[:2]
                    left_interval = lp.Interval(0, w/2*1.05, axis='x').put_on_canvas(img)
                    left_blocks = text_blocks.filter_by(left_interval, center=True)
                    left_blocks.sort(key = lambda b:b.coordinates[1], inplace=True)
                # The b.coordinates[1] corresponds to the y coordinate of the region
                # sort based on that can simulate the top-to-bottom reading order
                    right_blocks = lp.Layout([b for b in text_blocks if b not in left_blocks])
                    right_blocks.sort(key = lambda b:b.coordinates[1], inplace=True)

                # And finally combine the two lists and add the index
                    text_blocks = lp.Layout([b.set(id = idx) for idx, b in enumerate(left_blocks + right_blocks)])
                # Specify the path where you want to save the Word document

                    ocr_agent = lp.TesseractAgent(languages='eng')
                    for idx,block in enumerate(text_blocks):
                        #print(block.type)
                        if  block.type == "List":
                            segment_image = (block
                                .pad(left=14, right=10, top=5, bottom=5)
                                .crop_image(img))

                            # Preprocess the segment image
                            preprocessed_image = preprocess_image(segment_image)
                            #img_path = f"preprocessed_image_{idx}.png"
                            #cv2.imwrite(img_path, preprocessed_image)

                            # Add padding in each image segment can help improve robustness

                            text = pytess_for_list(preprocessed_image)
                            text_to_add = sanitize_text(text)
                            block.set(text=text_to_add, inplace=True)

                        # Add text to the Word document
                            add_list_to_docx(document, text) 
                        elif block.type == "Text":
                            segment_image = (block
                                .pad(left=10, right=10, top=5, bottom=5)
                                .crop_image(img))

                            # Preprocess the segment image
                            preprocessed_image = preprocess_image(segment_image)
                            #img_path = f"preprocessed_image_{idx}.png"
                            #cv2.imwrite(img_path, preprocessed_image)

                            # Add padding in each image segment can help improve robustness

                            text = pytess_for_list(preprocessed_image)
                            text_to_add = sanitize_text(text)+'\n'
                            block.set(text=text_to_add, inplace=True)

                        # Add text to the Word document
                            add_text_to_docx(document, text) 
                        elif block.type == "Title":
                            segment_image = (block
                                .pad(left=10, right=10, top=5, bottom=5)
                                .crop_image(img))

                           
                            # Add padding in each image segment can help improve robustness

                            text = ocr_agent.detect(segment_image)
                            text_to_add = sanitize_text(text)+'\n'
                            block.set(text=text_to_add, inplace=True)

                            # Add text to the Word document
                            add_title_to_docx(document, text_to_add)
                        elif block.type == "Figure":
                            segment_image = (block
                                .pad(left=10, right=10, top=2, bottom=2)
                                .crop_image(img))

                            # Preprocess the segment image
                            preprocessed_image = preprocess_image(segment_image)

                            # Save the preprocessed image with a unique name
                            if not isinstance(preprocessed_image, np.ndarray):
                                preprocessed_image = np.array(preprocessed_image)
                            img_path = "figure_{}.png".format(idx)
                            cv2.imwrite(img_path, preprocessed_image)

                            # Add the figure to the Word document
                            add_figure_to_docx(document, img_path)
                        elif block.type == "Table":
                            #print("inside table")
                            segment_image = (block
                                .pad(left=10, right=10, top=2, bottom=15)
                                .crop_image(img))

                            # Preprocess the segment image
                            preprocessed_image = preprocess_image(segment_image)
                            if not isinstance(preprocessed_image, np.ndarray):
                                preprocessed_image = np.array(preprocessed_image)
                            # The above code is performing the following tasks:
                            img_path = "table{}.png".format(idx)
                            cv2.imwrite(img_path, preprocessed_image)
                            add_figure_table_to_docx(document, img_path)
                            TD_THRESHOLD = 0.40
                            TSR_THRESHOLD = 0.40
                            padd_top, padd_left, padd_bottom, padd_right = 25, 10, 25, 5
                            delta_xmin, delta_ymin, delta_xmax, delta_ymax = 0,0,10,10
                            expand_rowcol_bbox_top, expand_rowcol_bbox_bottom = 5,5
                            table_pipeline.start_process(document,img_path, TD_THRESHOLD, TSR_THRESHOLD, padd_top, padd_left, padd_bottom, padd_right, delta_xmin, delta_ymin, delta_xmax, delta_ymax, expand_rowcol_bbox_top, expand_rowcol_bbox_bottom)

            document.save(docx_path)
            images_directory = "images"
            shutil.rmtree(images_directory)

            # Remove the 'output' directory
            output_directory = "output"
            shutil.rmtree(output_directory)

            # Define the extensions to remove
            extensions = [".jpg", ".png", ".pdf"]

            # List all files in the current directory
            current_directory = os.getcwd()
            all_files = os.listdir(current_directory)
                # Remove files with specified extensions
            for file in all_files:
                if any(file.endswith(ext) for ext in extensions):
                    os.remove(file)
            return FileResponse(
                docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="output_scanned.docx",
                status_code=200,
            )
        else:
            docx_file = 'output_unscanned.docx'
            # convert pdf to docx
            parse(pdf_file.filename, docx_file)
            return FileResponse(docx_file, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="output_unscanned.docx")
    except Exception as e:
        # Handle any other exceptions during processing
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error processing PDF: {str(e)}")    
def pdf_to_text(pdf):
    pages = pdf2image.convert_from_bytes(pdf, poppler_path="poppler/bin")
    ocr = PaddleOCR(use_angle_cls=True, lang='en')
    text_results = []
    for page_image in pages:
        img_np = np.array(page_image)
        result = ocr.ocr(img_np, cls=True)
        result = result[0]  # Assuming the first result is sufficient
        text = '\n'.join([line[1][0] for line in result])  # Join text lines into a single string
        text_results.append(text)
    return text_results

@app.post("/upload-file/")
async def upload_file(pdf_file: UploadFile):
    if not os.path.exists("output"):
        os.makedirs("output", exist_ok=True)
    
    # Read PDF file content
    pdf_content = await pdf_file.read()
    # Convert PDF to images
    text_results = pdf_to_text(pdf_content)
    print(text_results)
    # Perform OCR on images
    return {"message": text_results}


if __name__ == "__main__":
    uvicorn.run(app, host='0.0.0.0', port=8000)